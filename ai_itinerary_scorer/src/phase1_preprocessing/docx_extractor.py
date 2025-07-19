"""
src/phase1_preprocessing/docx_extractor.py

Extracts high-quality, context-rich text from .docx itinerary documents.
Preserves logical order, headings, and list structure for downstream NLP.
"""

from docx import Document

def extract_docx_text(file_path: str) -> str:
    """
    Extracts structured text (headings, paragraphs, lists) from a DOCX file.
    
    Returns:
        str: The document text with basic structure (days, activities, notes).
    """
    doc = Document(file_path)
    lines = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        
        style = para.style.name.lower()
        
        # Headings (e.g., "Day 1")
        if style.startswith('heading'):
            lines.append(f"\n# {text}\n")
        # Bullet/Numbered lists
        elif style in ('list bullet', 'list bullet 2', 'list number', 'list number 2', 'bullet', 'numbered list'):
            lines.append(f"- {text}")
        # Normal paragraph
        else:
            lines.append(text)

    # Optionally: Extract text from tables if present (very rare in itineraries)
    # Uncomment below if needed.
    #
    # for i, table in enumerate(doc.tables):
    #     lines.append(f"\n[Table {i+1}]\n")
    #     for row in table.rows:
    #         row_text = " | ".join(cell.text.strip() for cell in row.cells)
    #         lines.append(row_text)

    return "\n".join(lines)


# For debugging/demo usage
if __name__ == "__main__":
    import sys
    if len(sys.argv) != 2:
        print("Usage: python docx_extractor.py <docx_file>")
        exit(1)
    fp = sys.argv[1]
    print(extract_docx_text(fp))
